from docx import Document
from docx.shared import Inches
import os, glob, tkinter as tk
from tkinter import filedialog
import openpyxl
from openpyxl_image_loader import SheetImageLoader
import io

root = tk.Tk()
root.withdraw()

# This allows selection of excel file to be transferred from

file  = openpyxl.load_workbook(filedialog.askopenfilename(title = "Select the excel file to be transferred from."), data_only=True)

# This returns the sheet names of the excel file

sheet = file.sheetnames

# This selects the sheet that will be worked with

sheet = file[sheet[0]]

# This loads the images from the sheet into a dictionary

image_loader = SheetImageLoader(sheet)

# This selects the words document that everything will be transfered to

worder = filedialog.askopenfilename(title = "Select the word file to be transferred to.")

# This puts the word document in a form that python can use

doc = Document(worder)

# Selects the second tablein the document

table = doc.tables[1]

# This counter will help find where in the excel sheet I want to start

counter = 0

# This loop helps find the starting point for copying  from excel spread sheet which starts at step one

for row in sheet.iter_rows(1, 20, values_only= True):

	counter += 1


	if row[0] == 1:
		break

# This loop will check to see if the cell contains a number(step) or a string(merged cells)

for row in sheet.iter_rows(counter, sheet.max_row,  values_only= True):

	# Adds a new row to the table and saves the cells as a variable

	rowcells = table.add_row().cells

	if type(row[0]).__name__ == 'str':
		
		rowcells[0].paragraphs[0].add_run(row[0])

	else :

		# Takes the third cell in the word table and puts the text from the 3rd column of excel file into it. 

		rowcells[2].paragraphs[0].add_run(row[2])

		# Puts string None into the 4th cell in the row 

		rowcells[3].paragraphs[0].add_run('None')

		# Takes the fifth cell in the word table and puts the text from the 4th column of excel file into it. 

		rowcells[4].paragraphs[0].add_run(row[3])

		# Not every cell has a picture so try statement skips pasts if there is not picture

		try:

			# Gets the image from the row of the excel file that is being pulled from.

			# All pictures in column B in this file.

			image = image_loader.get("B" + str(counter))

			# Using this to save a photo stops from having to save photo to hard drive so it can be loaded into word file

			imageByte = io.BytesIO()

			# Saving the image

			image.save(imageByte, format= "BMP")

			# Adding the image to the proper place in the word table (2nd cell)

			run = rowcells[1].paragraphs[0].add_run()

			run.add_picture(imageByte, width = Inches(2), height = Inches(2))

		except :
			next


		
	# Counter keeps up with the row in excel that is being added in.
	

	counter +=1


		
# Saving the word document with all the changes. 

doc.save(worder)

	









